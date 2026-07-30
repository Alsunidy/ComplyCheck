"""
engine/evidence_pipeline.py

Multi-evidence ingestion pipeline for SAMA CSF audits. Turns a heterogeneous
evidence package (policies, procedures, screenshots, audit reports, Excel
matrices, config exports, ...) into provenance-tagged text chunks ready for
control mapping and assessment in compliance_engine.py.

Stages (the "agents" of the pipeline):
  1. classify_evidence()   - Document Classification: decide the evidence type
                             from the user's declared type or the file itself.
  2. extract_text()        - Text Extraction / OCR: per-type text extraction;
                             images go through the OpenAI vision model, which
                             OCRs visible text AND describes what the
                             screenshot demonstrates (Evidence Understanding
                             for images happens in the same call).
  3. build_evidence_chunks() - orchestrates 1+2 for a whole package and tags
                             every chunk with its source file and metadata.

Control mapping + assessment live in compliance_engine.run_evidence_compliance_check().
"""
import base64
import json
import os
from io import BytesIO
from pathlib import Path

from dotenv import load_dotenv

import arabic_text
import progress

load_dotenv(Path(__file__).resolve().parent.parent.parent / ".env")

# Optional preferred vision model; llm_client walks its fallback chain from
# here. Every model in the default chain supports image input.
VISION_MODEL = os.environ.get("COMPLYCHECK_VISION_MODEL")

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
SPREADSHEET_EXTENSIONS = {".xlsx", ".xlsm"}
DOCUMENT_EXTENSIONS = {".pdf", ".docx"}
TEXT_EXTENSIONS = {".txt", ".json", ".yaml", ".yml", ".cfg", ".ini", ".conf", ".xml", ".csv", ".log"}

SUPPORTED_EXTENSIONS = (
    IMAGE_EXTENSIONS | SPREADSHEET_EXTENSIONS | DOCUMENT_EXTENSIONS | TEXT_EXTENSIONS
)

# Evidence types the classifier can assign. Users may also declare one of
# these explicitly in the upload metadata, which always wins.
EVIDENCE_TYPES = [
    "policy",
    "procedure",
    "standard",
    "screenshot",
    "audit_report",
    "compliance_matrix",
    "configuration",
    "other",
]

VISION_PROMPT = (
    "You are analyzing a screenshot submitted as cyber security compliance "
    "evidence for a SAMA CSF audit. Do two things:\n"
    "1. OCR: transcribe all readable text in the image.\n"
    "2. Understanding: describe what system/setting the screenshot shows and "
    "what security practice it demonstrates (e.g. 'Azure AD portal showing "
    "MFA enabled for all users').\n"
    "Respond with a JSON object with exactly these keys:\n"
    '  "ocr_text": string (the transcribed text, empty string if none)\n'
    '  "description": string (1-3 sentences on what the screenshot demonstrates)\n'
    "Do not invent content that is not visible in the image."
)


# ---------------------------------------------------------------------------
# Stage 1: Document Classification Agent
# ---------------------------------------------------------------------------
_FILENAME_TYPE_HINTS = [
    ("policy", "policy"),
    ("procedure", "procedure"),
    ("standard", "standard"),
    ("audit", "audit_report"),
    ("matrix", "compliance_matrix"),
    ("config", "configuration"),
]


def classify_evidence(filename: str, declared_type: str | None = None) -> str:
    """Return the evidence type, preferring the user's declared type."""
    if declared_type and declared_type in EVIDENCE_TYPES:
        return declared_type

    extension = Path(filename).suffix.lower()
    if extension in IMAGE_EXTENSIONS:
        return "screenshot"
    if extension in SPREADSHEET_EXTENSIONS:
        return "compliance_matrix"
    if extension in TEXT_EXTENSIONS:
        return "configuration"

    name_lower = filename.lower()
    for hint, ev_type in _FILENAME_TYPE_HINTS:
        if hint in name_lower:
            return ev_type
    return "policy" if extension in DOCUMENT_EXTENSIONS else "other"


# ---------------------------------------------------------------------------
# Stage 2: Text Extraction / OCR Agent
# ---------------------------------------------------------------------------
def _read_pdf(content: bytes) -> str:
    import fitz  # pymupdf

    with fitz.open(stream=content, filetype="pdf") as doc:
        return "\n".join(page.get_text() for page in doc)


def _read_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _read_spreadsheet(content: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                lines.append(" | ".join(cells))
    workbook.close()
    return "\n".join(lines)


def _read_text_file(content: bytes) -> str:
    return content.decode("utf-8", errors="replace")


def _analyze_image_with_vision(content: bytes, filename: str) -> str:
    """OCR + understand a screenshot via the OpenAI vision model.

    Returns combined text: the model's description of what the screenshot
    demonstrates, followed by the OCR'd text. Raises RuntimeError if the
    vision call keeps failing so the caller can surface a clear error.
    """
    import llm_client

    extension = Path(filename).suffix.lower().lstrip(".")
    media_type = "jpeg" if extension in ("jpg", "jpeg") else "png"
    image_b64 = base64.b64encode(content).decode()

    parts = [
        {"type": "text", "text": VISION_PROMPT},
        {
            "type": "image_url",
            "image_url": {"url": f"data:image/{media_type};base64,{image_b64}"},
        },
    ]

    last_error = None
    for attempt in range(2):
        try:
            text = llm_client.generate_json(
                contents=parts,
                preferred_model=VISION_MODEL,
            )
            data = json.loads(text)
            description = str(data.get("description", "")).strip()
            ocr_text = str(data.get("ocr_text", "")).strip()
            combined = "\n".join(part for part in [f"Screenshot analysis: {description}", ocr_text] if part)
            if not combined.strip():
                raise ValueError("Vision model returned empty analysis.")
            return combined
        except RuntimeError as exc:
            raise RuntimeError(f"Vision analysis failed for '{filename}': {exc}") from exc
        except (ValueError, json.JSONDecodeError) as exc:
            last_error = exc
    raise RuntimeError(f"Vision analysis for '{filename}' returned unusable output: {last_error}")


def _chunk_spreadsheet_rows(text: str) -> list[str]:
    """One chunk per data row, carrying its sheet name as context.

    A whole compliance matrix as a single chunk mixes unrelated topics (e.g.
    training, vendor risk, access recertification) into one blob, which then
    sits at a moderate, non-specific similarity to *every* control instead of
    a sharp match to the one row that's actually relevant. Per-row chunks
    keep each topic isolated so retrieval can tell them apart.
    """
    chunks: list[str] = []
    current_sheet = ""
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("[Sheet:"):
            current_sheet = line.strip("[]")
            continue
        chunks.append(f"{current_sheet}: {line}" if current_sheet else line)
    return chunks


def extract_text(content: bytes, filename: str) -> str:
    """Extract raw text from any supported evidence file.

    Arabic output is repaired for the lam-alef ligature defect that PDF
    extractors introduce (see arabic_text), so Arabic evidence embeds and
    displays correctly.
    """
    extension = Path(filename).suffix.lower()
    if extension == ".pdf":
        return _read_pdf(content)
    if extension == ".docx":
        return _read_docx(content)
    if extension == ".doc":
        raise ValueError(
            "Legacy .doc files are not supported. Please re-save as .docx or PDF."
        )
    if extension in SPREADSHEET_EXTENSIONS:
        return _read_spreadsheet(content)
    if extension in IMAGE_EXTENSIONS:
        return _analyze_image_with_vision(content, filename)
    if extension in TEXT_EXTENSIONS:
        return _read_text_file(content)
    raise ValueError(f"Unsupported file type: {extension}")


def _extract_and_clean(content: bytes, filename: str) -> str:
    text = extract_text(content, filename)
    if arabic_text.detect_language(text) == "ar":
        text = arabic_text.clean_arabic(text)
    return text


# ---------------------------------------------------------------------------
# Stage 3: package orchestration -> provenance-tagged chunks
# ---------------------------------------------------------------------------
def build_evidence_chunks(files: list[dict], chunk_text_fn) -> tuple[list[dict], list[dict], str]:
    """Process a whole evidence package.

    `files`: list of {"filename", "content" (bytes), "evidence_type" (optional),
             "category" (optional), "description" (optional),
             "related_control" (optional)}.
    `chunk_text_fn`: text -> list[str] chunker (compliance_engine.chunk_text).

    Returns (chunks, manifest, language):
      chunks:   [{"text", "source", "evidence_type", "category",
                  "description", "related_control"}]
      manifest: per-file metadata summary for the report's evidence_used list.
      language: "ar" or "en", detected across the package's combined text --
                decides which SAMA corpus is matched against and which
                language the report and answers are written in.
    """
    chunks: list[dict] = []
    manifest: list[dict] = []

    progress.update(stage="classifying", total=len(files), current=0)
    for item in files:
        filename = item["filename"]
        evidence_type = classify_evidence(filename, item.get("evidence_type"))
        progress.update(
            stage="analyzing_images" if evidence_type == "screenshot" else "extracting",
            detail=filename,
        )
        text = _extract_and_clean(item["content"], filename)
        extension = Path(filename).suffix.lower()

        if evidence_type == "screenshot":
            # A screenshot's whole analysis is one small unit of evidence;
            # splitting it would separate the description from the OCR text.
            file_chunks = [text]
        elif extension in SPREADSHEET_EXTENSIONS:
            file_chunks = _chunk_spreadsheet_rows(text)
        else:
            file_chunks = chunk_text_fn(text)

        if not file_chunks:
            raise ValueError(f"No readable content found in '{filename}'.")

        # Prefix the user's description so retrieval can match on it too.
        description = (item.get("description") or "").strip()
        for chunk in file_chunks:
            chunk_body = f"{description}\n{chunk}".strip() if description else chunk
            chunks.append(
                {
                    "text": chunk_body,
                    "source": filename,
                    "evidence_type": evidence_type,
                    "category": (item.get("category") or "").strip(),
                    "description": description,
                    "related_control": (item.get("related_control") or "").strip(),
                }
            )

        manifest.append(
            {
                "filename": filename,
                "type": evidence_type,
                "category": (item.get("category") or "").strip(),
                "description": description,
                "related_control": (item.get("related_control") or "").strip(),
                "chunk_count": len(file_chunks),
            }
        )

        progress.advance()

    # Detect once over the whole package: a single mixed-language file
    # shouldn't flip the report language on its own.
    progress.update(stage="detecting_language", detail="")
    language = arabic_text.detect_language(" ".join(c["text"] for c in chunks))
    return chunks, manifest, language
